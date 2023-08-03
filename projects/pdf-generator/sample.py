from docx.shared import Inches, Pt
from docx import Document

def gen_report(frame=None, num=None):
        """
        Generate PV report for storage purpose

        Data:
            model       : Model number
            jabil_sn    : Jabil serial number
            test_date   : test date
            failure     : issues
            ms          : operator name
            time_taken  : time used for PV

        Example:
        job_dict[report] = {
                                'model': 'xxxxx-xxxxx',
                                'jabil_sn': 'MYxxxxxxxx',
                                'test_date': '15th June 2023',
                                'failure': 'led/scratches',
                                'ms': 'JS Peh',
                                'time_taken': '30 mins'
                                }
        """
        if frame:
                image_name = str(job) + '_' + str(num) if num else str(job)
                debug_imshow(frame, save='report/' + image_name)
        else:
            job_dict['report']['test_date'] = datetime.datetime.now()
            doc = Document('template.docx')
            t_in = time.time()
            # Loop through paragraphs and replace placeholders with data from dictionary
            for paragraph in doc.paragraphs:
                # Reconstruct paragraph content
                new_paragraph_text = paragraph.text
                for key, value in job_dict['report']:
                    placeholder = '{{' + key + '}}'
                    if placeholder in new_paragraph_text:
                        new_paragraph_text = new_paragraph_text.replace(placeholder, value)
                # Clear existing runs
                for run in paragraph.runs:
                    run.text = ''
                # Add the new text with formatting
                run = paragraph.add_run(new_paragraph_text)
                font = run.font
                font.name = 'Times New Roman'
                font.size = Pt(12)

            for paragraph in doc.paragraphs:
                for placeholder, name in job_dict['report']['image'].items():
                    if placeholder in paragraph.text:
                        # Clear placeholder
                        for run in paragraph.runs:
                            # Construct image path and insert image
                            image_path = os.path.join('report/', name)
                            paragraph.add_run().add_picture(image_path, width=Inches(6))

            # Save as docx
            doc.save('output.docx')

            # Step 4: Convert Word to PDF
            word = comtypes.client.CreateObject('Word.Application')
            # Using raw string format
            doc_com = word.Documents.Open(r'C:??\report.docx')
            # Specify the absolute path for output.pdf as well, or it will be saved in an unintended location
            doc_com.SaveAs(r'C:??\output.pdf', FileFormat=17)  # 17 represents pdf format
            doc_com.Close()
            word.Quit()
